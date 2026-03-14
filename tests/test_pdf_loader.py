"""Tests for PDF loading and chunking."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from docfinder.ingestion.pdf_loader import build_chunks, get_pdf_metadata, iter_text_parts
from docfinder.models import ChunkRecord


class TestIterTextParts:
    """Test iter_text_parts function."""

    @patch("docfinder.ingestion.pdf_loader.fitz")
    def test_iter_text_parts_simple(self, mock_fitz: MagicMock, tmp_path: Path) -> None:
        """Should extract text from PDF page by page."""
        # Setup mock
        mock_page = MagicMock()
        mock_page.get_text.return_value = "Page 1 text"
        mock_page.find_tables.return_value = MagicMock(tables=[])

        mock_doc = MagicMock()
        mock_doc.__len__ = MagicMock(return_value=1)
        mock_doc.__getitem__ = MagicMock(return_value=mock_page)

        mock_fitz.open.return_value = mock_doc

        # Test
        pdf_path = tmp_path / "test.pdf"
        pdf_path.write_bytes(b"dummy pdf content")

        parts = list(iter_text_parts(pdf_path))

        assert len(parts) == 1
        assert "Page 1 text" in parts[0]

    @patch("docfinder.ingestion.pdf_loader.fitz")
    def test_iter_text_parts_multiple_pages(self, mock_fitz: MagicMock, tmp_path: Path) -> None:
        """Should extract text from multiple pages."""
        # Setup mock pages
        mock_page1 = MagicMock()
        mock_page1.get_text.return_value = "Page 1"
        mock_page1.find_tables.return_value = MagicMock(tables=[])
        mock_page2 = MagicMock()
        mock_page2.get_text.return_value = "Page 2"
        mock_page2.find_tables.return_value = MagicMock(tables=[])
        mock_page3 = MagicMock()
        mock_page3.get_text.return_value = "Page 3"
        mock_page3.find_tables.return_value = MagicMock(tables=[])

        pages = [mock_page1, mock_page2, mock_page3]

        mock_doc = MagicMock()
        mock_doc.__len__ = MagicMock(return_value=3)
        mock_doc.__getitem__ = MagicMock(side_effect=lambda i: pages[i])

        mock_fitz.open.return_value = mock_doc

        pdf_path = tmp_path / "multi.pdf"
        pdf_path.write_bytes(b"dummy")

        parts = list(iter_text_parts(pdf_path))

        assert len(parts) == 3
        assert "Page 1" in parts[0]
        assert "Page 2" in parts[1]
        assert "Page 3" in parts[2]

    @patch("docfinder.ingestion.pdf_loader.fitz")
    @patch("docfinder.ingestion.pdf_loader.LOGGER")
    def test_iter_text_parts_page_error(
        self, mock_logger: MagicMock, mock_fitz: MagicMock, tmp_path: Path
    ) -> None:
        """Should log warning and continue on page extraction error."""
        mock_page1 = MagicMock()
        mock_page1.get_text.return_value = "Page 1"
        mock_page1.find_tables.return_value = MagicMock(tables=[])

        mock_page2 = MagicMock()
        mock_page2.find_tables.return_value = MagicMock(tables=[])
        mock_page2.get_text.side_effect = Exception("Extraction failed")

        mock_page3 = MagicMock()
        mock_page3.get_text.return_value = "Page 3"
        mock_page3.find_tables.return_value = MagicMock(tables=[])

        pages = [mock_page1, mock_page2, mock_page3]

        mock_doc = MagicMock()
        mock_doc.__len__ = MagicMock(return_value=3)
        mock_doc.__getitem__ = MagicMock(side_effect=lambda i: pages[i])

        mock_fitz.open.return_value = mock_doc

        pdf_path = tmp_path / "error.pdf"
        pdf_path.write_bytes(b"dummy")

        parts = list(iter_text_parts(pdf_path))

        # Should extract text from pages that work
        assert len(parts) == 2
        assert "Page 1" in parts[0]
        assert "Page 3" in parts[1]
        # Should have logged warning
        assert mock_logger.warning.called

    @patch("docfinder.ingestion.pdf_loader.fitz")
    @patch("docfinder.ingestion.pdf_loader.LOGGER")
    def test_iter_text_parts_open_error(
        self, mock_logger: MagicMock, mock_fitz: MagicMock, tmp_path: Path
    ) -> None:
        """Should log error and return empty on file open failure."""
        mock_fitz.open.side_effect = Exception("Cannot open file")

        pdf_path = tmp_path / "broken.pdf"
        pdf_path.write_bytes(b"dummy")

        parts = list(iter_text_parts(pdf_path))

        assert len(parts) == 0
        assert mock_logger.error.called


class TestGetPdfMetadata:
    """Test get_pdf_metadata function."""

    @patch("docfinder.ingestion.pdf_loader.fitz")
    def test_get_metadata_simple(self, mock_fitz: MagicMock, tmp_path: Path) -> None:
        """Should extract metadata from PDF."""
        mock_doc = MagicMock()
        mock_doc.__len__ = MagicMock(return_value=1)
        mock_doc.metadata = {"title": "Test PDF"}

        mock_fitz.open.return_value = mock_doc

        pdf_path = tmp_path / "test.pdf"
        pdf_path.write_bytes(b"dummy")

        metadata = get_pdf_metadata(pdf_path)

        assert metadata["title"] == "Test PDF"
        assert metadata["page_count"] == "1"

    @patch("docfinder.ingestion.pdf_loader.fitz")
    def test_get_metadata_none(self, mock_fitz: MagicMock, tmp_path: Path) -> None:
        """Should handle PDF without metadata."""
        mock_doc = MagicMock()
        mock_doc.__len__ = MagicMock(return_value=1)
        mock_doc.metadata = None

        mock_fitz.open.return_value = mock_doc

        pdf_path = tmp_path / "no_meta.pdf"
        pdf_path.write_bytes(b"dummy")

        metadata = get_pdf_metadata(pdf_path)

        assert metadata["title"] == "no_meta"
        assert metadata["page_count"] == "1"

    @patch("docfinder.ingestion.pdf_loader.fitz")
    def test_get_metadata_empty_title(self, mock_fitz: MagicMock, tmp_path: Path) -> None:
        """Should use filename when title is empty."""
        mock_doc = MagicMock()
        mock_doc.__len__ = MagicMock(return_value=5)
        mock_doc.metadata = {"title": ""}

        mock_fitz.open.return_value = mock_doc

        pdf_path = tmp_path / "empty_title.pdf"
        pdf_path.write_bytes(b"dummy")

        metadata = get_pdf_metadata(pdf_path)

        assert metadata["title"] == "empty_title"
        assert metadata["page_count"] == "5"


class TestBuildChunks:
    """Test build_chunks function."""

    @patch("docfinder.ingestion.pdf_loader.iter_text_parts_paged")
    @patch("docfinder.ingestion.pdf_loader.get_pdf_metadata")
    def test_build_chunks_simple(
        self, mock_meta: MagicMock, mock_iter: MagicMock, tmp_path: Path
    ) -> None:
        """Should build chunks from PDF text."""
        mock_meta.return_value = {"title": "Test", "page_count": "1"}
        mock_iter.return_value = iter([(1, "This is a test document with some content.")])

        pdf_path = tmp_path / "test.pdf"
        pdf_path.write_bytes(b"dummy")

        chunks = list(build_chunks(pdf_path, max_chars=50, overlap=10))

        assert len(chunks) >= 1
        assert all(isinstance(chunk, ChunkRecord) for chunk in chunks)
        assert chunks[0].document_path == pdf_path
        assert chunks[0].index == 0
        assert chunks[0].metadata["title"] == "Test"

    @patch("docfinder.ingestion.pdf_loader.iter_text_parts_paged")
    @patch("docfinder.ingestion.pdf_loader.get_pdf_metadata")
    def test_build_chunks_multiple(
        self, mock_meta: MagicMock, mock_iter: MagicMock, tmp_path: Path
    ) -> None:
        """Should create multiple chunks for long text."""
        # Use text with sentence boundaries so the sentence-aware chunker can split it
        long_text = " ".join(f"Sentence number {i} is here." for i in range(50))
        mock_meta.return_value = {"title": "Long", "page_count": "1"}
        mock_iter.return_value = iter([(1, long_text)])

        pdf_path = tmp_path / "long.pdf"
        pdf_path.write_bytes(b"dummy")

        chunks = list(build_chunks(pdf_path, max_chars=100, overlap=20))

        # Should create multiple chunks
        assert len(chunks) > 1
        # Chunks should be indexed sequentially
        for i, chunk in enumerate(chunks):
            assert chunk.index == i

    @patch("docfinder.ingestion.pdf_loader.iter_text_parts_paged")
    @patch("docfinder.ingestion.pdf_loader.get_pdf_metadata")
    def test_build_chunks_metadata(
        self, mock_meta: MagicMock, mock_iter: MagicMock, tmp_path: Path
    ) -> None:
        """Should include metadata in chunks."""
        mock_meta.return_value = {"title": "My Document", "page_count": "5"}
        mock_iter.return_value = iter([(1, "Content")])

        pdf_path = tmp_path / "doc.pdf"
        pdf_path.write_bytes(b"dummy")

        chunks = list(build_chunks(pdf_path))

        assert len(chunks) >= 1
        chunk = chunks[0]
        assert chunk.metadata["title"] == "My Document"

    @patch("docfinder.ingestion.pdf_loader.iter_text_parts_paged")
    @patch("docfinder.ingestion.pdf_loader.get_pdf_metadata")
    def test_build_chunks_empty_text(
        self, mock_meta: MagicMock, mock_iter: MagicMock, tmp_path: Path
    ) -> None:
        """Should handle PDF with no extractable text."""
        mock_meta.return_value = {"title": "Empty", "page_count": "1"}
        mock_iter.return_value = iter([])

        pdf_path = tmp_path / "empty.pdf"
        pdf_path.write_bytes(b"dummy")

        chunks = list(build_chunks(pdf_path))

        # Should return no chunks for empty text
        assert len(chunks) == 0


class TestTableExtraction:
    """Test PDF table extraction via _extract_page_text."""

    def test_table_to_markdown(self) -> None:
        """Should convert a table to Markdown format."""
        from docfinder.ingestion.pdf_loader import _table_to_markdown

        mock_table = MagicMock()
        mock_table.extract.return_value = [
            ["Name", "Age", "City"],
            ["Marco", "30", "Roma"],
            ["Laura", "25", "Milano"],
        ]
        md = _table_to_markdown(mock_table)
        assert "| Name | Age | City |" in md
        assert "| --- | --- | --- |" in md
        assert "| Marco | 30 | Roma |" in md
        assert "| Laura | 25 | Milano |" in md

    def test_table_to_markdown_empty(self) -> None:
        """Should return empty string for empty table."""
        from docfinder.ingestion.pdf_loader import _table_to_markdown

        mock_table = MagicMock()
        mock_table.extract.return_value = []
        assert _table_to_markdown(mock_table) == ""

    def test_table_to_markdown_none_cells(self) -> None:
        """Should handle None cell values gracefully."""
        from docfinder.ingestion.pdf_loader import _table_to_markdown

        mock_table = MagicMock()
        mock_table.extract.return_value = [
            ["Col1", None],
            [None, "value"],
        ]
        md = _table_to_markdown(mock_table)
        assert "| Col1 |  |" in md
        assert "|  | value |" in md

    @patch("docfinder.ingestion.pdf_loader.fitz")
    def test_extract_page_with_table(self, mock_fitz: MagicMock) -> None:
        """Should include Markdown table in extracted text."""
        from docfinder.ingestion.pdf_loader import _extract_page_text

        # Mock a table
        mock_table = MagicMock()
        mock_table.bbox = (0, 100, 500, 200)
        mock_table.extract.return_value = [
            ["Product", "Price"],
            ["Widget", "10.00"],
        ]

        mock_tables_result = MagicMock()
        mock_tables_result.tables = [mock_table]

        # Mock fitz.Rect to return proper rect objects
        mock_fitz.Rect.side_effect = lambda coords: MagicMock(
            y0=coords[1],
            width=coords[2] - coords[0],
            height=coords[3] - coords[1],
            is_empty=False,
            __and__=lambda self, other: MagicMock(
                is_empty=False, width=self.width, height=self.height
            ),
        )

        mock_page = MagicMock()
        mock_page.find_tables.return_value = mock_tables_result
        # Text blocks outside the table area
        mock_page.get_text.return_value = [
            (0, 10, 500, 50, "Introduction paragraph.", 0, 0),
        ]

        text = _extract_page_text(mock_page)
        assert "Product" in text
        assert "Price" in text
        assert "Widget" in text

    @patch("docfinder.ingestion.pdf_loader.fitz")
    def test_extract_page_no_tables(self, mock_fitz: MagicMock) -> None:
        """Should fall back to plain text when no tables found."""
        from docfinder.ingestion.pdf_loader import _extract_page_text

        mock_page = MagicMock()
        mock_page.find_tables.return_value = MagicMock(tables=[])
        mock_page.get_text.return_value = "Just plain text here."

        text = _extract_page_text(mock_page)
        assert text == "Just plain text here."

    def test_extract_page_find_tables_exception(self) -> None:
        """Should fall back to plain text when find_tables() raises."""
        from docfinder.ingestion.pdf_loader import _extract_page_text

        mock_page = MagicMock()
        mock_page.find_tables.side_effect = RuntimeError("unsupported")
        mock_page.get_text.return_value = "Fallback text."

        text = _extract_page_text(mock_page)
        assert text == "Fallback text."

    @patch("docfinder.ingestion.pdf_loader.fitz")
    def test_extract_page_skips_image_blocks(self, mock_fitz: MagicMock) -> None:
        """Should skip image blocks (block_type != 0)."""
        from docfinder.ingestion.pdf_loader import _extract_page_text

        mock_table = MagicMock()
        mock_table.bbox = (0, 200, 500, 300)
        mock_table.extract.return_value = [["A"], ["1"]]

        mock_tables_result = MagicMock()
        mock_tables_result.tables = [mock_table]

        mock_fitz.Rect.side_effect = lambda coords: MagicMock(
            y0=coords[1],
            width=coords[2] - coords[0],
            height=coords[3] - coords[1],
            is_empty=False,
            __and__=lambda self, other: MagicMock(is_empty=True),
        )

        mock_page = MagicMock()
        mock_page.find_tables.return_value = mock_tables_result
        mock_page.get_text.return_value = [
            (0, 10, 500, 50, "Text block.", 0, 0),  # text block — kept
            (0, 60, 500, 100, "image data", 1, 1),  # image block — skipped
        ]

        text = _extract_page_text(mock_page)
        assert "Text block." in text
        assert "image data" not in text

    @patch("docfinder.ingestion.pdf_loader.fitz")
    def test_extract_page_excludes_overlapping_text(self, mock_fitz: MagicMock) -> None:
        """Text blocks overlapping >50% with a table should be excluded."""
        from docfinder.ingestion.pdf_loader import _extract_page_text

        mock_table = MagicMock()
        mock_table.bbox = (0, 100, 500, 200)
        mock_table.extract.return_value = [["Col"], ["Val"]]

        mock_tables_result = MagicMock()
        mock_tables_result.tables = [mock_table]

        # Table rect
        table_rect = MagicMock(
            y0=100,
            width=500,
            height=100,
            is_empty=False,
        )
        # Block inside table: full overlap
        block_inside_rect = MagicMock(
            y0=110,
            width=500,
            height=30,
            is_empty=False,
        )
        # Block outside table: no overlap
        block_outside_rect = MagicMock(
            y0=10,
            width=500,
            height=40,
            is_empty=False,
        )

        rects = [table_rect, block_outside_rect, block_inside_rect]
        mock_fitz.Rect.side_effect = rects

        # Overlap intersection for inside block: full overlap → ratio = 1.0
        full_overlap = MagicMock(is_empty=False, width=500, height=30)
        block_inside_rect.__and__ = MagicMock(return_value=full_overlap)
        # No overlap for outside block
        no_overlap = MagicMock(is_empty=True)
        block_outside_rect.__and__ = MagicMock(return_value=no_overlap)

        mock_page = MagicMock()
        mock_page.find_tables.return_value = mock_tables_result
        mock_page.get_text.return_value = [
            (0, 10, 500, 50, "Outside text.", 0, 0),
            (0, 110, 500, 140, "Inside table text.", 1, 0),
        ]

        text = _extract_page_text(mock_page)
        assert "Outside text." in text
        assert "Inside table text." not in text

    def test_table_to_markdown_newlines_in_cells(self) -> None:
        """Should replace newlines in cell values with spaces."""
        from docfinder.ingestion.pdf_loader import _table_to_markdown

        mock_table = MagicMock()
        mock_table.extract.return_value = [
            ["Header"],
            ["line1\nline2\nline3"],
        ]
        md = _table_to_markdown(mock_table)
        assert "line1 line2 line3" in md
        assert "\n" not in md.split("\n")[2]  # data row has no embedded newlines

    def test_table_to_markdown_header_only(self) -> None:
        """Should handle a table with only a header row."""
        from docfinder.ingestion.pdf_loader import _table_to_markdown

        mock_table = MagicMock()
        mock_table.extract.return_value = [
            ["Col1", "Col2"],
        ]
        md = _table_to_markdown(mock_table)
        assert "| Col1 | Col2 |" in md
        assert "| --- | --- |" in md
        # Only 2 lines: header + separator
        assert len(md.strip().split("\n")) == 2


class TestIterTextPartsTxt:
    """Test plain text extraction."""

    def test_reads_text_file(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import iter_text_parts_txt

        f = tmp_path / "hello.txt"
        f.write_text("Hello world!", encoding="utf-8")
        parts = list(iter_text_parts_txt(f))
        assert len(parts) == 1
        assert "Hello world!" in parts[0]

    def test_read_error(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import iter_text_parts_txt

        f = tmp_path / "missing.txt"
        parts = list(iter_text_parts_txt(f))
        assert parts == []


class TestIterTextPartsTxtPaged:
    """Test plain text virtual paging — error path."""

    def test_read_error(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import iter_text_parts_txt_paged

        f = tmp_path / "missing.txt"
        pages = list(iter_text_parts_txt_paged(f))
        assert pages == []

    def test_whitespace_only_file(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import iter_text_parts_txt_paged

        f = tmp_path / "blank.txt"
        f.write_text("   \n\n  \n")
        pages = list(iter_text_parts_txt_paged(f))
        assert pages == []


class TestCleanMd:
    """Test Markdown formatting cleanup."""

    def test_strips_bold_italic(self) -> None:
        from docfinder.ingestion.pdf_loader import _clean_md

        assert "hello" in _clean_md("**hello**")
        assert "world" in _clean_md("*world*")
        assert "***" not in _clean_md("***bold italic***")

    def test_strips_links_keeps_text(self) -> None:
        from docfinder.ingestion.pdf_loader import _clean_md

        result = _clean_md("[click here](http://example.com)")
        assert "click here" in result
        assert "http://" not in result

    def test_strips_images(self) -> None:
        from docfinder.ingestion.pdf_loader import _clean_md

        result = _clean_md("![alt text](image.png)")
        assert "![" not in result
        assert "image.png" not in result

    def test_strips_code(self) -> None:
        from docfinder.ingestion.pdf_loader import _clean_md

        result = _clean_md("Use `print()` here")
        assert "`" not in result

    def test_strips_headings(self) -> None:
        from docfinder.ingestion.pdf_loader import _clean_md

        result = _clean_md("## My Title\nSome text")
        assert "##" not in result
        assert "My Title" in result

    def test_strips_horizontal_rule(self) -> None:
        from docfinder.ingestion.pdf_loader import _clean_md

        result = _clean_md("Above\n---\nBelow")
        assert "---" not in result
        assert "Above" in result
        assert "Below" in result


class TestIterTextPartsMd:
    """Test Markdown text extraction."""

    def test_extracts_cleaned_text(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import iter_text_parts_md

        f = tmp_path / "doc.md"
        f.write_text("# Title\nSome **bold** text.")
        parts = list(iter_text_parts_md(f))
        assert len(parts) == 1
        assert "bold" in parts[0]
        assert "**" not in parts[0]

    def test_read_error(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import iter_text_parts_md

        f = tmp_path / "missing.md"
        parts = list(iter_text_parts_md(f))
        assert parts == []


class TestIterTextPartsMdPaged:
    """Test Markdown paged extraction — error path."""

    def test_read_error(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import iter_text_parts_md_paged

        f = tmp_path / "missing.md"
        pages = list(iter_text_parts_md_paged(f))
        assert pages == []

    def test_empty_sections_skipped(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import iter_text_parts_md_paged

        f = tmp_path / "sparse.md"
        f.write_text("# Title\n\n\n# Another\nContent here.")
        pages = list(iter_text_parts_md_paged(f))
        # Only sections with actual content after cleaning
        assert all(text.strip() for _, text in pages)


class TestIterTextPartsDocx:
    """Test Word document extraction."""

    def test_extracts_paragraphs(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import iter_text_parts_docx

        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        f = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("")  # empty — should be skipped
        doc.add_paragraph("Second paragraph")
        doc.save(str(f))

        parts = list(iter_text_parts_docx(f))
        assert len(parts) == 2
        assert "First paragraph" in parts[0]
        assert "Second paragraph" in parts[1]

    def test_read_error(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import iter_text_parts_docx

        try:
            from docx import Document  # noqa: F401
        except ImportError:
            pytest.skip("python-docx not installed")

        f = tmp_path / "corrupt.docx"
        f.write_bytes(b"not a real docx")
        parts = list(iter_text_parts_docx(f))
        assert parts == []


class TestIterTextPartsDocxPaged:
    """Test Word document paged extraction — error and edge cases."""

    def test_read_error(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import iter_text_parts_docx_paged

        try:
            from docx import Document  # noqa: F401
        except ImportError:
            pytest.skip("python-docx not installed")

        f = tmp_path / "corrupt.docx"
        f.write_bytes(b"not a real docx")
        pages = list(iter_text_parts_docx_paged(f))
        assert pages == []

    def test_fewer_than_page_size(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import iter_text_parts_docx_paged

        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        f = tmp_path / "small.docx"
        doc = Document()
        for i in range(3):
            doc.add_paragraph(f"Para {i}")
        doc.save(str(f))

        pages = list(iter_text_parts_docx_paged(f))
        assert len(pages) == 1
        assert pages[0][0] == 1


class TestImportDocxDocument:
    """Test python-docx lazy import."""

    def test_returns_document_class(self) -> None:
        from docfinder.ingestion.pdf_loader import _import_docx_document

        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        result = _import_docx_document()
        assert result is Document

    @patch.dict("sys.modules", {"docx": None})
    def test_returns_none_when_not_installed(self) -> None:
        from docfinder.ingestion.pdf_loader import _import_docx_document

        result = _import_docx_document()
        assert result is None


class TestIterTextBySuffix:
    """Test the file-type dispatcher."""

    def test_pdf_dispatch(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import _iter_text_by_suffix

        with patch("docfinder.ingestion.pdf_loader.iter_text_parts") as mock:
            mock.return_value = iter(["pdf text"])
            parts = list(_iter_text_by_suffix(tmp_path / "file.pdf"))
            assert parts == ["pdf text"]
            mock.assert_called_once()

    def test_txt_dispatch(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import _iter_text_by_suffix

        f = tmp_path / "file.txt"
        f.write_text("hello")
        parts = list(_iter_text_by_suffix(f))
        assert "hello" in parts[0]

    def test_md_dispatch(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import _iter_text_by_suffix

        f = tmp_path / "file.md"
        f.write_text("content")
        parts = list(_iter_text_by_suffix(f))
        assert "content" in parts[0]

    def test_docx_dispatch(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import _iter_text_by_suffix

        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        f = tmp_path / "file.docx"
        doc = Document()
        doc.add_paragraph("docx content")
        doc.save(str(f))
        parts = list(_iter_text_by_suffix(f))
        assert any("docx content" in p for p in parts)

    def test_unsupported_extension(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import _iter_text_by_suffix

        f = tmp_path / "file.xyz"
        f.write_text("data")
        parts = list(_iter_text_by_suffix(f))
        assert parts == []


class TestGetTitle:
    """Test _get_title helper."""

    @patch("docfinder.ingestion.pdf_loader.get_pdf_metadata")
    def test_pdf_with_title(self, mock_meta: MagicMock, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import _get_title

        mock_meta.return_value = {"title": "My PDF Title"}
        result = _get_title(tmp_path / "doc.pdf")
        assert result == "My PDF Title"

    @patch("docfinder.ingestion.pdf_loader.get_pdf_metadata")
    def test_pdf_metadata_error(self, mock_meta: MagicMock, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import _get_title

        mock_meta.side_effect = Exception("cannot open")
        result = _get_title(tmp_path / "broken.pdf")
        assert result == "broken"

    def test_non_pdf_uses_stem(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import _get_title

        result = _get_title(tmp_path / "report.txt")
        assert result == "report"


class TestIterPagedText:
    """Test the paged text dispatcher."""

    def test_dispatches_pdf(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import _iter_paged_text

        with patch("docfinder.ingestion.pdf_loader.iter_text_parts_paged") as mock:
            mock.return_value = iter([(1, "page text")])
            pages = list(_iter_paged_text(tmp_path / "file.pdf"))
            assert pages == [(1, "page text")]

    def test_dispatches_md(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import _iter_paged_text

        f = tmp_path / "file.md"
        f.write_text("# Hello\nWorld")
        pages = list(_iter_paged_text(f))
        assert len(pages) >= 1

    def test_dispatches_txt(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import _iter_paged_text

        f = tmp_path / "file.txt"
        f.write_text("Some content")
        pages = list(_iter_paged_text(f))
        assert len(pages) == 1

    def test_dispatches_docx(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import _iter_paged_text

        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        f = tmp_path / "file.docx"
        doc = Document()
        doc.add_paragraph("Hello")
        doc.save(str(f))
        pages = list(_iter_paged_text(f))
        assert len(pages) >= 1

    def test_unsupported_returns_empty(self, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import _iter_paged_text

        f = tmp_path / "file.xyz"
        f.write_text("data")
        pages = list(_iter_paged_text(f))
        assert pages == []


class TestIterTextPartsPaged:
    """Test iter_text_parts_paged (PDF-specific paged extraction)."""

    @patch("docfinder.ingestion.pdf_loader.fitz")
    def test_yields_page_numbers(self, mock_fitz: MagicMock, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import iter_text_parts_paged

        mock_page1 = MagicMock()
        mock_page1.get_text.return_value = "First page"
        mock_page1.find_tables.return_value = MagicMock(tables=[])
        mock_page2 = MagicMock()
        mock_page2.get_text.return_value = "Second page"
        mock_page2.find_tables.return_value = MagicMock(tables=[])

        mock_doc = MagicMock()
        mock_doc.__len__ = MagicMock(return_value=2)
        mock_doc.__getitem__ = MagicMock(side_effect=[mock_page1, mock_page2])
        mock_fitz.open.return_value = mock_doc

        pages = list(iter_text_parts_paged(tmp_path / "test.pdf"))
        assert len(pages) == 2
        assert pages[0][0] == 1
        assert pages[1][0] == 2
        assert "First page" in pages[0][1]
        assert "Second page" in pages[1][1]

    @patch("docfinder.ingestion.pdf_loader.fitz")
    def test_open_error(self, mock_fitz: MagicMock, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import iter_text_parts_paged

        mock_fitz.open.side_effect = Exception("cannot open")
        pages = list(iter_text_parts_paged(tmp_path / "bad.pdf"))
        assert pages == []

    @patch("docfinder.ingestion.pdf_loader.fitz")
    def test_skips_empty_pages(self, mock_fitz: MagicMock, tmp_path: Path) -> None:
        from docfinder.ingestion.pdf_loader import iter_text_parts_paged

        mock_page1 = MagicMock()
        mock_page1.get_text.return_value = "Content"
        mock_page1.find_tables.return_value = MagicMock(tables=[])
        mock_page2 = MagicMock()
        mock_page2.get_text.return_value = ""
        mock_page2.find_tables.return_value = MagicMock(tables=[])

        mock_doc = MagicMock()
        mock_doc.__len__ = MagicMock(return_value=2)
        mock_doc.__getitem__ = MagicMock(side_effect=[mock_page1, mock_page2])
        mock_fitz.open.return_value = mock_doc

        pages = list(iter_text_parts_paged(tmp_path / "test.pdf"))
        assert len(pages) == 1
